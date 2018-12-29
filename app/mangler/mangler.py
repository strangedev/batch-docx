import os
import re
from typing import Dict

import datetime
import docx
from flask import json
from shutil import copyfile

import util
from bean.manglermapping import ManglerMapping
from mangler import template_lang
from mangler.databatch import Databatch


class Mangler(object):
    __pattern_placeholder = r"{([^{}]+)}"
    __matcher = re.compile(__pattern_placeholder)

    def __init__(self, template):

        self.__template = template  # type: template.Template
        self.__mapping_template = None  # type: Dict[str, str]
        self.__doc = docx.Document(template.path)  # type: docx.Document

    def _iterate_runs(self, doc=None):

        doc = self.__doc if doc is None else doc

        for p in doc.paragraphs:
            for r in p.runs:
                yield r

    @property
    def mapping_template(self) -> Dict[str, str]:

        if self.__mapping_template is None:

            self.__mapping_template = dict({})

            for r in self._iterate_runs():
                for m in Mangler.__matcher.findall(r.text):
                    if m == "%date%": continue
                    self.__mapping_template[m] = ""

        return self.__mapping_template

    def create(self, mapping: ManglerMapping, batch: Databatch):

        outfiles = []
        mappings = json.loads(mapping.mappings_json)  # type: Dict[str, str]

        for i in range(batch.row_count):

            phone_number = batch.get("Telefon Büro", i)
            outfile_path = util.generate_outfile_path(phone_number)

            copyfile(self.__template.path, outfile_path)
            doc = docx.Document(outfile_path)

            for run in self._iterate_runs(doc=doc):
                if "{%date%}" in run.text:
                    run.text = run.text.replace("{%date%}", datetime.date.today().strftime("%d.%m.%Y"))

            bindings = {template_attr: batch.get(batch_attr, i) for template_attr, batch_attr in mappings.items()}

            for template_attr, batch_attr in mappings.items():
                value = batch.get(batch_attr, i)

                for search_term, replace_term in json.loads(mapping.replacements_json).items():
                    # replace term might be template lang expression
                    replace_term = template_lang.eval_template_lang(replace_term, bindings)

                    if search_term.startswith("regex:"):
                        pattern_str = search_term.split(":", 1)[1]
                        try:
                            m = re.compile(pattern_str)
                            value = m.sub(replace_term, value, 1)
                        except:
                            value = "Regex fehlerhaft!"

                    elif search_term in value:
                        value = value.replace(search_term, replace_term, 1)

                to_replace = "{" + template_attr + "}"

                for run in self._iterate_runs(doc=doc):
                    if to_replace not in run.text:
                        continue

                    run.text = run.text.replace(to_replace, value)

            doc.save(outfile_path)
            outfiles.append(outfile_path)

        return outfiles
